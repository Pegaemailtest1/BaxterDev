import logging
import sys
import os
import base64
import re
import io
from collections import Counter, defaultdict
from docx import Document
from langchain.schema import Document as LangChainDoc
import pdfplumber
import fitz  # PyMuPDF
import numpy as np
from sklearn.cluster import KMeans
from PIL import Image, ImageOps, ImageEnhance
import pytesseract

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)

def preprocess_image_for_ocr(image):
    gray = image.convert("L")
    enhancer = ImageEnhance.Contrast(gray)
    contrast = enhancer.enhance(2.5)
    bw = contrast.point(lambda x: 0 if x < 180 else 255, '1')
    resized = bw.resize((bw.width * 2, bw.height * 2), Image.LANCZOS)
    return resized

def crop_top_right_quarter(image):
    width, height = image.size
    return image.crop((width * 0.5, 0, width, height * 0.5))

def extract_dates(text):
    matches = re.findall(r"(Issued Date|Effective Date)\s*[:\-]?\s*(\d{1,2}-[A-Za-z]{3}-\d{4})", text)
    return {label: date for label, date in matches}

def extract_issued_effective_date_from_page(page, page_index):
    
    # Step 1: Date string matcher, no label needed
    date_value_regex = re.compile(r"\d{1,2}[-/][A-Za-z]{3}[-/]\d{4}")

    # Step 2: Optional label prefix
    labeled_date_regex = re.compile(
        r"(Issued|Effective)\s*Date\s*[:\-]?\s*(\d{1,2}[-/][A-Za-z]{3}[-/]\d{4})",
        re.IGNORECASE
    )

    rect = page.rect
    crop_rect = fitz.Rect(rect.width * 0.6, 0, rect.width, rect.height * 0.07)
    pix = page.get_pixmap(dpi=300, clip=crop_rect, alpha=False)
    image = Image.open(io.BytesIO(pix.tobytes("png")))

    image = image.convert("L")
    image = ImageEnhance.Contrast(image).enhance(2.5)
    image = image.point(lambda x: 0 if x < 160 else 255, mode='1')

    ocr_text = pytesseract.image_to_string(image, config="--oem 3 --psm 11").strip()

    # First try labeled dates
    issued = None
    effective = None

    for label, date_str in labeled_date_regex.findall(ocr_text):
        if "issued" in label.lower():
            issued = date_str
        elif "effective" in label.lower():
            effective = date_str

    # If any date is still missing, fallback to positional date guessing
    if not issued or not effective:
        all_dates = date_value_regex.findall(ocr_text)

        if not issued and len(all_dates) > 0:
            issued = all_dates[0]
        if not effective and len(all_dates) > 1:
            effective = all_dates[1]

    if issued or effective:
        return {
            "page": page_index + 1,
            "issued_date": issued,
            "effective_date": effective,
            "ocr_text": ocr_text
        }

    return None



def is_multi_column_paragraph_page(page, min_word_threshold=30, min_cluster_distance=100, balance_threshold=0.25):
    words = page.extract_words()
    if len(words) < min_word_threshold:
        return False

    tables = page.extract_tables()
    table_word_count = sum(len(cell.split()) for table in tables for row in table for cell in row if cell)
    total_word_count = sum(len(w["text"].split()) for w in words)

    if table_word_count / (total_word_count + 1e-6) > 0.5:
        return False

    x_positions = np.array([[word['x0']] for word in words])
    kmeans = KMeans(n_clusters=2, n_init=10, random_state=42)
    labels = kmeans.fit_predict(x_positions)
    cluster_sizes = Counter(labels)
    centers = sorted([center[0] for center in kmeans.cluster_centers_])

    balance = min(cluster_sizes.values()) / sum(cluster_sizes.values())
    center_distance = abs(centers[1] - centers[0])

    return balance > balance_threshold and center_distance > min_cluster_distance

def looks_like_section_title(line):
    line = line.strip()
    if not line:
        return False

    words = line.split()
    if len(words) <= 12:
        capitalized = sum(1 for w in words if w[0].isupper() or w.isupper())
        if capitalized / len(words) >= 0.6:
            return True

    if re.match(r"^[A-Z][\w\s\[\],\-()]{3,}$", line):
        return True

    if line.isupper() and len(line.split()) <= 10:
        return True

    return False
def normalize_section_name(section_title):
    # Lowercase
    title = section_title.lower()
    # Remove numbering and punctuation
    title = re.sub(r'^[\d\.\-\)\(a-zA-Z\s]*', '', title)  # Remove prefixes like '3.2.', 'Section III:', etc.
    title = re.sub(r'[^\w\s]', '', title)  # Remove punctuation
    # Replace spaces and hyphens with underscore
    title = re.sub(r'[\s\-]+', '_', title.strip())
    return title

def extract_header_footer_lines(page, height, margin_ratio=0.1, line_spacing=5):
    top_limit = height * margin_ratio
    bottom_limit = height * (1 - margin_ratio)
    words = page.extract_words()

    def group_words_by_line(words, y_thresh=5):
        lines = []
        current_line = []
        last_y = None

        for word in sorted(words, key=lambda w: (w["top"], w["x0"])):
            y = word["top"]
            if last_y is None or abs(y - last_y) <= y_thresh:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(current_line)
                current_line = [word]
            last_y = y

        if current_line:
            lines.append(current_line)

        return lines

    all_lines = group_words_by_line(words, y_thresh=line_spacing)

    header_lines = []
    footer_lines = []

    for line_words in all_lines:
        if not line_words:
            continue
        y_top = line_words[0]["top"]
        line_text = " ".join(w["text"] for w in line_words).strip()
        if y_top <= top_limit:
            header_lines.append(line_text)
        elif y_top >= bottom_limit:
            footer_lines.append(line_text)

    return header_lines, footer_lines

def raster_image_to_svg_path(image_bytes, threshold=200):
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert("L")
        binary = img.point(lambda x: 255 if x > threshold else 0, mode='1')
        pixels = binary.load()

        path = []
        for y in range(binary.height):
            for x in range(binary.width):
                if pixels[x, y] == 0:
                    path.append(f"M{x},{y}h1v1h-1z")
        return " ".join(path) if path else None
    except Exception:
        return None

def clean_table_and_remove_empty_cells(raw_table: str) -> str:
    cleaned_rows = []

    for line in raw_table.strip().split("\n"):
        cells = [cell.strip() for cell in line.split("|")]
        
        # Remove empty cells
        non_empty_cells = [cell for cell in cells if cell and cell != ""]

        # Only keep rows with at least 2 columns of meaningful content
        if len(non_empty_cells) >= 2:
            cleaned_line = " | ".join(non_empty_cells)
            cleaned_rows.append(cleaned_line)

    return "\n".join(cleaned_rows)

def extract_pdf_text_tables_images(pdf_path):
    extracted_content = []
    header_footer_map = {"headers": [], "footers": []}
    unique_headers = set()
    unique_footers = set()

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            width = page.width
            height = page.height

            header_lines, footer_lines = extract_header_footer_lines(page, height)
            header_hash = " | ".join(header_lines).strip()
            footer_hash = " | ".join(footer_lines).strip()

            if header_hash and header_hash not in unique_headers:
                unique_headers.add(header_hash)
                header_footer_map["headers"].append(header_hash)

            if footer_hash and footer_hash not in unique_footers:
                unique_footers.add(footer_hash)
                header_footer_map["footers"].append(footer_hash)

            full_text = page.extract_text() or ""
            text_lines = full_text.split("\n")
            filtered_lines = [
                line for line in text_lines
                if line.strip() not in header_lines and line.strip() not in footer_lines
            ]
            combined_text = "\n".join(filtered_lines)

            if "IFU_LABEL" in pdf_path and is_multi_column_paragraph_page(page):
                column_width = width / 3
                column_texts = []
                for i in range(3):
                    left = i * column_width
                    right = (i + 1) * column_width
                    column = page.within_bbox((left, 0, right, height))
                    col_text = column.extract_text() or ""
                    column_texts.append(col_text.strip())
                combined_text = "\n\n".join(f"[Column {i+1}]\n{txt}" for i, txt in enumerate(column_texts) if txt.strip())
            else:
                combined_text = page.extract_text() or ""

            section_chunks = defaultdict(list)
            current_section = "Introduction"
            lines = combined_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if looks_like_section_title(line):
                    current_section = line.rstrip(":").strip()
                    continue
                section_chunks[current_section].append(line)

            for section_title, content_lines in section_chunks.items():
                if not content_lines:
                    continue
                section_text = f"[{section_title}]\n" + "\n".join(content_lines).strip()
                normalized_section = normalize_section_name(section_title)
                if section_text:
                    extracted_content.append({
                        "type": "section-text",
                        "section": section_title,
                        "normalized_section": normalized_section,
                        "page": page_num + 1,
                        "content": section_text
                    })

            tables = page.extract_tables()
            for idx, table in enumerate(tables):
                cleaned_rows = [" | ".join((cell or "").replace("\n", " ").strip() for cell in row) for row in table]
                table_str = "\n".join(cleaned_rows)

                if table_str.strip():
                    cleaned_table_str = clean_table_and_remove_empty_cells(table_str)
                    extracted_content.append({
                        "type": "table",
                        "section": f"Table {idx+1}",
                        "normalized_section": f"table_{idx+1}",
                        "page": page_num + 1,
                        "content": cleaned_table_str
                    })


            if combined_text.strip():
                extracted_content.append({
                    "type": "section-table",
                    "section": "Tables",
                    "page": page_num + 1,
                    "normalized_section": "tables",
                    "content": combined_text.strip()
                })

    pdf_doc = fitz.open(pdf_path)
    header_height_ratio = 0.1

    for page_index in range(len(pdf_doc)):
        page = pdf_doc[page_index]

        #top banner extraction
        top__banner_content = extract_issued_effective_date_from_page(page, page_index)

        if top__banner_content:  # Already checks if something was found
            issued = top__banner_content.get("issued_date") or "N/A"
            effective = top__banner_content.get("effective_date") or "N/A"
            ocr_preview = top__banner_content.get("ocr_text", "").strip()
            top__banner_page_num = top__banner_content["page"]

            content_str = f"Issued Date: {issued}, Effective Date: {effective}, OCR Text: {ocr_preview}"
            #logging.info(f"Extracted Header Date on Page {page_num}: {content_str}")
            if content_str.strip():
                extracted_content.append({
                    "type": "header-date",
                    "page": top__banner_page_num,
                    "section": "header-date",
                    "content": content_str
                })
            #logging.info(f"Extracted content:{extracted_content}")
        # Extract images from the page
        images = page.get_images(full=True)

        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            b64_image = base64.b64encode(image_bytes).decode("utf-8")

            image_rects = page.get_image_rects(xref)
            for rect in image_rects:
                is_header = rect.y0 < page.rect.height * header_height_ratio
                ocr_text = ""

                if is_header:
                    try:
                        ocr_image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                        ocr_image = crop_top_right_quarter(ocr_image)
                        ocr_image = preprocess_image_for_ocr(ocr_image)
                        custom_config = r'--oem 3 --psm 11'
                        ocr_text = pytesseract.image_to_string(ocr_image, lang="eng", config=custom_config).strip()
                        if ocr_text:
                            logging.info(f"OCR Header Text on Page {page_index+1}, Image {img_index+1}: {ocr_text[:100]}...")
                        else:
                            logging.warning(f"No OCR text found in header image Page {page_index+1}, Image {img_index+1}")
                    except Exception as ocr_error:
                        logging.warning(f"OCR failed on Page {page_index+1}, Header Image {img_index+1}: {ocr_error}")
                        ocr_text = ""

                vector_path = raster_image_to_svg_path(image_bytes) if is_header else None

                extracted_content.append({
                    "type": "header-image" if is_header else "image",
                    "page": page_index + 1,
                    "format": image_ext,
                    "content": f"[Image {img_index+1} on Page {page_index+1}]\nOCR Text:\n{ocr_text}",
                    "image_base64": b64_image,
                    "vector_path": vector_path
                })


    if header_footer_map["headers"]:
        extracted_content.insert(0, {
            "type": "header",
            "page": 0,
            "content": "[Headers]\n" + "\n".join(header_footer_map["headers"])
        })
    if header_footer_map["footers"]:
        extracted_content.append({
            "type": "footer",
            "page": -1,
            "content": "[Footers]\n" + "\n".join(header_footer_map["footers"])
        })
    
    return extracted_content

def extract_docx_text_tables_images(docx_path):
    extracted_content = []
    doc = Document(docx_path)
    extracted_text = "\n".join([p.text for p in doc.paragraphs])

    rows_text = []
    for table_index, table in enumerate(doc.tables):
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                rows_text.append(f"Table {table_index + 1}: {row_text}")

    image_data = []
    rels = doc.part._rels
    for rel in rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            image_ext = rel.target_part.content_type.split("/")[-1]
            b64_image = base64.b64encode(image_bytes).decode('utf-8')
            image_data.append(f"[Word Image ({image_ext}) base64]: {b64_image[:100]}...")

    if rows_text:
        extracted_text += "\n\n[TABLES]\n" + "\n".join(rows_text)
    if image_data:
        extracted_text += "\n\n[IMAGES]\n" + "\n".join(image_data)

    return extracted_text

def load_documents(folder_path):
    all_docs = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.startswith("~$"):
                continue

            full_path = os.path.join(root, file)
            logging.info(f"Processing file: {full_path}")

            try:
                document_id = os.path.splitext(os.path.basename(full_path))[0]
                if file.lower().endswith(".pdf"):
                    full_text = extract_pdf_text_tables_images(full_path)
                    for entry in full_text:
                        all_docs.append(LangChainDoc(
                            page_content=entry["content"],                            
                            metadata={
                                "source": full_path,
                                "document_id": document_id,
                                "type": entry["type"],
                                "page": entry["page"],
                                **({"section": entry["section"]} if "section" in entry else {}),
                                **({"normalized_section": entry["normalized_section"]} if "normalized_section" in entry else {}),
                            }
                        ))

                elif file.lower().endswith((".doc", ".docx")):
                    full_text = extract_docx_text_tables_images(full_path)
                    all_docs.append(LangChainDoc(
                        page_content=full_text,
                        metadata={"source": full_path, "document_id": document_id,"type": "docx"}
                    ))
            except Exception as e:
                logging.error(f"Failed to process {full_path} - {e}")
                continue

    return all_docs