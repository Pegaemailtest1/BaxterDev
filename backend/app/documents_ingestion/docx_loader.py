import logging
import sys


import logging
from docx import Document
import base64
from langchain.schema import Document as LangChainDoc

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)

def extract_docx_text_tables_images(docx_path):
    
    extracted_content = []    
    doc = Document(docx_path)

    # --- Extract text from paragraphs ---
    extracted_content = "\n".join([p.text for p in doc.paragraphs])

    # --- Extract tables ---
    rows_text = []
    for table_index, table in enumerate(doc.tables):
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                rows_text.append(f"Table {table_index + 1}: {row_text}")

    # --- Extract embedded images ---
    image_data = []
    rels = doc.part._rels
    for rel in rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            image_ext = rel.target_part.content_type.split("/")[-1]  # e.g., 'jpeg', 'png'
            b64_image = base64.b64encode(image_bytes).decode('utf-8')
            image_data.append(f"[Word Image ({image_ext}) base64]: {b64_image[:100]}...")  # truncated preview

    # --- Combine extracted content ---
    if rows_text:
        extracted_content += "\n\n[TABLES]\n" + "\n".join(rows_text)
    if image_data:
        extracted_content += "\n\n[IMAGES]\n" + "\n".join(image_data)

    return extracted_content

